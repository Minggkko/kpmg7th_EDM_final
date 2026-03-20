"""
report_exporter.py
─────────────────────────────────────────────────────────────────────────────
목적:
    사용자가 CONFIRM 후 선택한 포맷으로 보고서를 내보냅니다.
    지원 포맷: PDF / DOCX / HWP

포맷별 구현 방식:
    PDF  : reportlab Platypus (Python 순수 구현, 외부 의존 없음)
    DOCX : python-docx (Python 순수 구현, 외부 의존 없음)
    HWP  : DOCX 생성 후 LibreOffice 헤드리스 변환
           → LibreOffice가 설치되어 있어야 합니다.
           → 미설치 시 HWPExportError 발생 (graceful)

HWP 변환 전제조건:
    - LibreOffice 설치 필요
    - Windows  : C:/Program Files/LibreOffice/program/soffice.exe
    - macOS    : /Applications/LibreOffice.app/.../soffice
    - Linux    : soffice (PATH에 존재)

사용 예:
    from report_editor   import load_draft
    from report_exporter import export_report, ExportFormat

    draft = load_draft("esg_report_output_draft.json")
    path  = export_report(draft, output_dir="exports", fmt=ExportFormat.PDF)
    path  = export_report(draft, output_dir="exports", fmt=ExportFormat.DOCX)
    path  = export_report(draft, output_dir="exports", fmt=ExportFormat.HWP)

프론트엔드 연동:
    - 사용자가 포맷 선택 후 CONFIRM → export_report() 호출
    - 반환된 파일 절대경로를 프론트에 전달 → 다운로드 트리거
"""

from __future__ import annotations

import os
import shutil
import subprocess
import tempfile
from datetime import datetime
from enum import Enum
from pathlib import Path


# ── 지원 포맷 ──────────────────────────────────────────────────────────────────

class ExportFormat(str, Enum):
    PDF  = "pdf"
    DOCX = "docx"
    HWP  = "hwp"


class HWPExportError(RuntimeError):
    """LibreOffice 미설치 등 HWP 변환 불가 시 발생합니다."""


# ── LibreOffice 실행 파일 탐색 경로 ───────────────────────────────────────────

_LIBREOFFICE_CANDIDATES = [
    r"C:\Program Files\LibreOffice\program\soffice.exe",
    r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    "/Applications/LibreOffice.app/Contents/MacOS/soffice",
    "soffice",          # Linux PATH
    "libreoffice",      # Linux 일부 배포판
]


# ═══════════════════════════════════════════════════════════════════════════════
# 공개 인터페이스
# ═══════════════════════════════════════════════════════════════════════════════

def export_report(
    draft: dict,
    output_dir: str = "exports",
    fmt: ExportFormat | str = ExportFormat.PDF,
) -> str:
    """
    초안을 지정된 포맷 파일로 내보냅니다.

    Args:
        draft      : load_draft()로 로드한 초안 dict.
                     편집된 current 값을 기준으로 내보냅니다.
        output_dir : 저장 디렉터리 (없으면 자동 생성).
        fmt        : ExportFormat 또는 문자열 "pdf" / "docx" / "hwp".

    Returns:
        생성된 파일의 절대 경로.

    Raises:
        ValueError     : 지원하지 않는 포맷.
        HWPExportError : LibreOffice 미설치 등 HWP 변환 실패.
    """
    fmt = ExportFormat(fmt) if isinstance(fmt, str) else fmt

    Path(output_dir).mkdir(parents=True, exist_ok=True)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    base_name = f"esg_report_{timestamp}"

    if fmt == ExportFormat.PDF:
        out_path = os.path.join(output_dir, f"{base_name}.pdf")
        _export_pdf(draft, out_path)

    elif fmt == ExportFormat.DOCX:
        out_path = os.path.join(output_dir, f"{base_name}.docx")
        _export_docx(draft, out_path)

    elif fmt == ExportFormat.HWP:
        out_path = os.path.join(output_dir, f"{base_name}.hwp")
        _export_hwp(draft, out_path)

    else:
        raise ValueError(f"지원하지 않는 포맷: {fmt}")

    abs_path = os.path.abspath(out_path)
    print(f"  [Exporter] {fmt.value.upper()} 저장: {abs_path}")
    return abs_path


# ═══════════════════════════════════════════════════════════════════════════════
# PDF 내보내기 (reportlab Platypus)
# ═══════════════════════════════════════════════════════════════════════════════

def _export_pdf(draft: dict, output_path: str) -> None:
    """reportlab Platypus로 PDF를 생성합니다. 한글은 맑은 고딕을 사용합니다."""
    from reportlab.lib           import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles   import ParagraphStyle
    from reportlab.lib.units    import mm
    from reportlab.platypus     import (
        SimpleDocTemplate, Paragraph, Spacer, Table,
        TableStyle, HRFlowable
    )

    font  = _register_korean_font()

    def S(name, **kw):
        return ParagraphStyle(name, fontName=font, **kw)

    style_h1 = S("H1", fontSize=18, spaceAfter=6,  leading=24, textColor=colors.HexColor("#1A365D"))
    style_h2 = S("H2", fontSize=13, spaceAfter=4,  leading=20, textColor=colors.HexColor("#2C5282"), spaceBefore=12)
    style_h3 = S("H3", fontSize=10, spaceAfter=3,  leading=15, textColor=colors.HexColor("#2D3748"), spaceBefore=6)
    style_p  = S("P",  fontSize=9,  spaceAfter=3,  leading=14, textColor=colors.HexColor("#4A5568"))
    style_cm = S("CM", fontSize=9,  spaceAfter=3,  leading=14, textColor=colors.HexColor("#2D3748"),
                 backColor=colors.HexColor("#EBF8FF"), borderPadding=5)
    style_nd = S("ND", fontSize=8,  spaceAfter=2,  leading=12, textColor=colors.HexColor("#A0AEC0"))

    story = []

    # 표지
    story.append(Paragraph("ESG 보고서", style_h1))
    story.append(Paragraph(
        f"생성일: {draft.get('generated_at', '')}  |  버전: {draft.get('version', 1)}",
        style_p
    ))
    story.append(Spacer(1, 6 * mm))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#CBD5E0")))
    story.append(Spacer(1, 4 * mm))

    for section in draft["sections"]:
        story.append(Paragraph(f"▶ {section['label']}", style_h2))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#E2E8F0")))

        for item in section["items"]:
            story.append(Paragraph(item["title"], style_h3))
            story.append(Paragraph(item["context"]["current"], style_p))

            for dp in item["data_points"]:
                if not dp["has_data"]:
                    story.append(Paragraph(
                        f"[{dp['indicator_code']}] {dp['dp_name']} — 입력된 데이터가 없습니다.",
                        style_nd
                    ))
                    continue

                story.append(Paragraph(
                    f"<b>[{dp['indicator_code']}] {dp['dp_name']}</b>  (단위: {dp['unit']})",
                    style_p
                ))
                story.append(_build_pdf_table(dp, font))
                story.append(Spacer(1, 2 * mm))

            story.append(Paragraph(
                f"<b>ESG 관점 해석</b><br/>{item['commentary']['current']}",
                style_cm
            ))
            story.append(Spacer(1, 4 * mm))

    doc = SimpleDocTemplate(
        output_path, pagesize=A4,
        leftMargin=20*mm, rightMargin=20*mm,
        topMargin=20*mm,  bottomMargin=20*mm,
    )
    doc.build(story)


def _build_pdf_table(dp: dict, font_name: str):
    """단일 data_point의 행들을 reportlab Table로 변환합니다."""
    from reportlab.lib        import colors
    from reportlab.lib.units  import mm
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.platypus   import Table, TableStyle, Paragraph

    cs = ParagraphStyle("Cell", fontName=font_name, fontSize=8, leading=12)

    header = [
        Paragraph("<b>보고기간</b>",               cs),
        Paragraph("<b>사이트</b>",                  cs),
        Paragraph(f"<b>측정값 ({dp['unit']})</b>",  cs),
    ]
    rows = [header]
    for r in dp["rows"]:
        val     = r.get("value")
        val_str = f"{val:,.3f}" if isinstance(val, (int, float)) else str(val or "-")
        rows.append([
            Paragraph((r.get("reporting_date") or "")[:7], cs),
            Paragraph(r.get("site_id") or "-",             cs),
            Paragraph(val_str,                             cs),
        ])

    tbl = Table(rows, colWidths=[35*mm, 55*mm, 45*mm], repeatRows=1)
    tbl.setStyle(TableStyle([
        ("BACKGROUND",     (0, 0), (-1, 0),  colors.HexColor("#2C5282")),
        ("TEXTCOLOR",      (0, 0), (-1, 0),  colors.white),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#EBF4FF")]),
        ("BOX",            (0, 0), (-1, -1), 0.5, colors.HexColor("#CBD5E0")),
        ("INNERGRID",      (0, 0), (-1, -1), 0.3, colors.HexColor("#E2E8F0")),
        ("TOPPADDING",     (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING",  (0, 0), (-1, -1), 3),
        ("LEFTPADDING",    (0, 0), (-1, -1), 4),
    ]))
    return tbl


def _register_korean_font() -> str:
    """맑은 고딕 등 한글 폰트를 탐색하여 reportlab에 등록합니다."""
    from reportlab.pdfbase         import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    candidates = [
        ("MalgunGothic", r"C:\Windows\Fonts\malgun.ttf"),
        ("NanumGothic",  r"C:\Windows\Fonts\NanumGothic.ttf"),
        ("AppleGothic",  "/Library/Fonts/AppleGothic.ttf"),
        ("NotoSansCJK",  "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"),
    ]
    for name, path in candidates:
        if os.path.exists(path):
            try:
                pdfmetrics.registerFont(TTFont(name, path))
                return name
            except Exception:
                continue

    return "Helvetica"   # 한글 미지원 fallback


# ═══════════════════════════════════════════════════════════════════════════════
# DOCX 내보내기 (python-docx)
# ═══════════════════════════════════════════════════════════════════════════════

def _export_docx(draft: dict, output_path: str) -> None:
    """python-docx로 .docx 파일을 생성합니다."""
    from docx                  import Document
    from docx.shared           import Pt, RGBColor, Cm
    from docx.enum.text        import WD_ALIGN_PARAGRAPH

    doc = Document()

    for sec in doc.sections:
        sec.top_margin    = Cm(2)
        sec.bottom_margin = Cm(2)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.5)

    # 표지
    title = doc.add_heading("ESG 보고서", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if title.runs:
        _set_run_color(title.runs[0], "1A365D")

    doc.add_paragraph(
        f"생성일: {draft.get('generated_at', '')}  |  버전: {draft.get('version', 1)}"
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    for section in draft["sections"]:
        h2 = doc.add_heading(f"▶ {section['label']}", level=1)
        if h2.runs:
            _set_run_color(h2.runs[0], "2C5282")

        for item in section["items"]:
            h3 = doc.add_heading(item["title"], level=2)
            if h3.runs:
                _set_run_color(h3.runs[0], "2D3748")

            doc.add_paragraph(item["context"]["current"])

            for dp in item["data_points"]:
                if not dp["has_data"]:
                    p   = doc.add_paragraph()
                    run = p.add_run(
                        f"[{dp['indicator_code']}] {dp['dp_name']} — 입력된 데이터가 없습니다."
                    )
                    run.font.color.rgb = RGBColor(0xA0, 0xAE, 0xC0)
                    run.font.size      = Pt(9)
                    continue

                label = doc.add_paragraph()
                run   = label.add_run(
                    f"[{dp['indicator_code']}] {dp['dp_name']}  (단위: {dp['unit']})"
                )
                run.bold      = True
                run.font.size = Pt(9)

                _add_docx_table(doc, dp)
                doc.add_paragraph()

            # AI 코멘터리
            cm_para = doc.add_paragraph()
            title_run = cm_para.add_run("ESG 관점 해석\n")
            title_run.bold      = True
            title_run.font.size = Pt(9)
            _set_run_color(title_run, "2C5282")
            body_run = cm_para.add_run(item["commentary"]["current"])
            body_run.font.size = Pt(9)
            _shade_paragraph(cm_para, "EBF8FF")
            doc.add_paragraph()

    doc.save(output_path)


def _add_docx_table(doc, dp: dict) -> None:
    """단일 data_point의 행들을 docx Table로 추가합니다."""
    from docx.shared import Pt, RGBColor

    rows  = dp["rows"]
    table = doc.add_table(rows=1 + len(rows), cols=3)
    table.style = "Table Grid"

    for i, text in enumerate(["보고기간", "사이트", f"측정값 ({dp['unit']})"]):
        cell = table.rows[0].cells[i]
        cell.text = text
        run  = cell.paragraphs[0].runs[0]
        run.bold           = True
        run.font.size      = Pt(8)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade_cell(cell, "2C5282")

    for row_idx, r in enumerate(rows):
        val     = r.get("value")
        val_str = f"{val:,.3f}" if isinstance(val, (int, float)) else str(val or "-")
        cells   = table.rows[row_idx + 1].cells
        for cell, text in zip(cells, [(r.get("reporting_date") or "")[:7], r.get("site_id") or "-", val_str]):
            cell.text = text
            run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(text)
            run.font.size = Pt(8)
            if row_idx % 2 == 1:
                _shade_cell(cell, "EBF4FF")


# ═══════════════════════════════════════════════════════════════════════════════
# HWP 내보내기 (DOCX → LibreOffice 변환)
#
# HWP(.hwp)는 한글과컴퓨터의 독점 바이너리 포맷으로
# Python에서 직접 생성하는 공식 라이브러리가 없습니다.
# LibreOffice 헤드리스 모드를 통해 DOCX → HWP 변환합니다.
#
# LibreOffice 설치 필요:
#   Windows: https://www.libreoffice.org/download/
#   macOS  : brew install --cask libreoffice
#   Linux  : sudo apt install libreoffice
# ═══════════════════════════════════════════════════════════════════════════════

def _export_hwp(draft: dict, output_path: str) -> None:
    """
    DOCX를 임시 생성한 뒤 LibreOffice로 HWP로 변환합니다.

    Raises:
        HWPExportError: LibreOffice가 없거나 변환에 실패한 경우.
    """
    soffice = _find_libreoffice()
    if soffice is None:
        raise HWPExportError(
            "HWP 내보내기에는 LibreOffice가 필요합니다.\n"
            "설치 후 재시도하세요: https://www.libreoffice.org/download/"
        )

    with tempfile.TemporaryDirectory() as tmp_dir:
        # 1단계: 임시 DOCX 생성
        tmp_docx = os.path.join(tmp_dir, "temp_report.docx")
        _export_docx(draft, tmp_docx)

        # 2단계: LibreOffice 헤드리스 변환 (DOCX → HWP)
        result = subprocess.run(
            [soffice, "--headless", "--convert-to", "hwp", "--outdir", tmp_dir, tmp_docx],
            capture_output=True,
            text=True,
            timeout=120,
        )

        if result.returncode != 0:
            raise HWPExportError(
                f"LibreOffice 변환 실패 (returncode={result.returncode})\n"
                f"stderr: {result.stderr.strip()}"
            )

        # 3단계: 변환된 HWP 파일을 최종 경로로 이동
        tmp_hwp = os.path.join(tmp_dir, "temp_report.hwp")
        if not os.path.exists(tmp_hwp):
            raise HWPExportError(
                "LibreOffice 변환은 성공했으나 HWP 파일을 찾을 수 없습니다.\n"
                "LibreOffice HWP 필터 플러그인 설치를 확인하세요."
            )
        shutil.move(tmp_hwp, output_path)


def _find_libreoffice() -> str | None:
    """시스템에서 LibreOffice 실행 파일 경로를 탐색합니다."""
    for candidate in _LIBREOFFICE_CANDIDATES:
        # 절대 경로인 경우 파일 존재 여부 확인
        if os.path.isabs(candidate):
            if os.path.exists(candidate):
                return candidate
        else:
            # PATH에서 탐색
            found = shutil.which(candidate)
            if found:
                return found
    return None


# ═══════════════════════════════════════════════════════════════════════════════
# 공통 유틸
# ═══════════════════════════════════════════════════════════════════════════════

def _set_run_color(run, hex_color: str) -> None:
    from docx.shared import RGBColor
    r, g, b = int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16)
    run.font.color.rgb = RGBColor(r, g, b)


def _shade_cell(cell, hex_color: str) -> None:
    from docx.oxml.ns import qn
    from docx.oxml    import OxmlElement
    tc_pr = cell._tc.get_or_add_tcPr()
    shd   = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tc_pr.append(shd)


def _shade_paragraph(para, hex_color: str) -> None:
    from docx.oxml.ns import qn
    from docx.oxml    import OxmlElement
    p_pr = para._p.get_or_add_pPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    p_pr.append(shd)
